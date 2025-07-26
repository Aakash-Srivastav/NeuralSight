from flask import Flask, request, jsonify, send_file, session
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
from flask_jwt_extended import JWTManager, create_access_token, jwt_required, get_jwt_identity
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime, timedelta
from PIL import Image
import numpy as np
import base64
import io
import os
import pandas as pd
from tensorflow.keras.models import load_model
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import atexit
import threading
import json

# Load environment variables
load_dotenv()

# Initialize Flask app
app = Flask(__name__)

# Configuration
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY')
app.config['JWT_SECRET_KEY'] = os.getenv('JWT_SECRET_KEY')
app.config['JWT_ACCESS_TOKEN_EXPIRES'] = timedelta(days=7)

# PostgreSQL database config
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('SQLALCHEMY_DATABASE_URI')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Initialize extensions
db = SQLAlchemy(app)
jwt = JWTManager(app)
CORS(app, resources={r"/api/*": {"origins": "http://localhost:3000"}})

# Global list to track temporary files for cleanup
temp_files_to_cleanup = []
cleanup_lock = threading.Lock()

# Global dictionary to store patient data per user session
user_patient_data = {}
patient_data_lock = threading.Lock()

# Database Model
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(100), unique=True, nullable=False)
    password_hash = db.Column(db.String(200), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    predictions = db.relationship('Prediction_table', backref='user', lazy=True)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

class Prediction_table(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    image_data = db.Column(db.LargeBinary, nullable=False)
    image_filename = db.Column(db.String(200), nullable=False)
    disease = db.Column(db.String(100), nullable=False)
    confidence = db.Column(db.Text, nullable=False)
    recommendations = db.Column(db.Text)
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)
    
    # Patient data from Excel
    patient_id = db.Column(db.String(50), nullable=True)
    patient_name = db.Column(db.String(200), nullable=True)
    patient_age = db.Column(db.Integer, nullable=True)
    patient_symptoms = db.Column(db.Text, nullable=True)


class DiseasePredictor:
    def __init__(self):
        self.model = load_model('models/new_cnn_model.keras')

    def preprocess(self, image_bytes):
        IMG_SIZE = (64, 64)
        img = Image.open(io.BytesIO(image_bytes)).convert('RGB')
        img = img.resize(IMG_SIZE)
        img_array = np.array(img).astype('float32')
        img_array = np.expand_dims(img_array, axis=0)
        return img_array

    def predict(self, image_bytes):
        img = self.preprocess(image_bytes)
        prediction = self.model.predict(img)
        predicted_class = np.argmax(prediction, axis=1)[0]
        confidence = np.max(self.model.predict(img)[0])

        diseases = ['Central Serous Chorioretinopathy-Color Fundus', 'Diabetic Retinopathy', 'Disc Edema',
                    'Glaucoma', 'Healthy', 'Macular Scar', 'Myopia', 'Pterygium', 'Retinal Detachment', 'Retinitis Pigmentosa']
        disease = diseases[predicted_class]

        # import random
        # confidence = random.uniform(0.7, 0.95)

        recommendations = {
            'Central Serous Chorioretinopathy-Color Fundus': 'Consult a doctor immediately.',
            'Diabetic Retinopathy': 'Consult a doctor immediately.',
            'Disc Edema': 'Consult a doctor immediately.',
            'Glaucoma': 'Consult a doctor immediately.',
            'Healthy': 'Continue maintaining good health habits.',
            'Macular Scar': 'Consult a doctor immediately.',
            'Myopia': 'Consult a doctor immediately.',
            'Pterygium': 'Self-isolate and seek medical attention.',
            'Retinal Detachment': 'Seek immediate medical attention.',
            'Retinitis Pigmentosa': 'Consult an ophthalmologist.'
        }

        return {
            'disease': disease,
            'confidence': str(confidence),
            'recommendations': recommendations.get(disease, 'Consult a healthcare professional.')
        }

predictor = DiseasePredictor()


def cleanup_temp_files():
    """Clean up temporary files"""
    with cleanup_lock:
        for temp_file_path in temp_files_to_cleanup[:]:
            try:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                temp_files_to_cleanup.remove(temp_file_path)
            except Exception as e:
                print(f"Error cleaning up temp file {temp_file_path}: {e}")

def schedule_cleanup(temp_file_path, delay=300):  # 5 minutes delay
    """Schedule cleanup of a temporary file"""
    def delayed_cleanup():
        import time
        time.sleep(delay)
        try:
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
            with cleanup_lock:
                if temp_file_path in temp_files_to_cleanup:
                    temp_files_to_cleanup.remove(temp_file_path)
        except Exception as e:
            print(f"Error in delayed cleanup of {temp_file_path}: {e}")
    
    cleanup_thread = threading.Thread(target=delayed_cleanup)
    cleanup_thread.daemon = True
    cleanup_thread.start()

def process_excel_file(excel_file):
    """Process uploaded Excel file and return patient data dictionary"""
    try:
        # Reset file pointer to beginning
        excel_file.seek(0)
        
        # Read Excel file
        df = pd.read_excel(excel_file)
        
        # Printing column names and first few rows for debugging (was not getting reponse)
        print(f"Excel columns: {df.columns.tolist()}")
        print(f"Excel shape: {df.shape}")
        print(f"First few rows:\n{df.head()}")
        
        # Convert column names to lowercase for consistent access
        df.columns = df.columns.str.lower().str.strip()
        
        # Creating dictionary with image_name as key
        patient_data = {}
        for index, row in df.iterrows():
            # Different possible column names for image
            image_name = None
            for col in ['image_name', 'image', 'filename', 'file_name', 'image_file']:
                if col in df.columns and pd.notna(row.get(col)):
                    image_name = str(row.get(col)).strip()
                    break
            
            if image_name and image_name.lower() != 'nan':
                patient_info = {
                    'id': str(row.get('id', '')) if pd.notna(row.get('id')) else '',
                    'full_name': str(row.get('full_name', '')) if pd.notna(row.get('full_name')) else '',
                    'age': int(float(row.get('age', 0))) if pd.notna(row.get('age')) and str(row.get('age')).replace('.', '').replace('-', '').isdigit() else None,
                    'symptoms': str(row.get('symptoms', '')) if pd.notna(row.get('symptoms')) else ''
                }
                patient_data[image_name] = patient_info
                print(f"Row {index}: Processed patient data for '{image_name}': {patient_info}")
        
        print(f"Total patient records processed: {len(patient_data)}")
        print(f"Patient data keys: {list(patient_data.keys())}")
        return patient_data
    except Exception as e:
        print(f"Error processing Excel file: {e}")
        import traceback
        traceback.print_exc()
        return {}

def find_patient_data(filename, patient_data):
    """Find patient data for a given filename with fuzzy matching"""
    print(f"Looking for patient data for filename: '{filename}'")
    print(f"Available patient data keys: {list(patient_data.keys())}")
    
    if not patient_data:
        print("No patient data available")
        return {}
    
    # Trying exact match first
    if filename in patient_data:
        print(f"Found exact match for '{filename}'")
        return patient_data[filename]
    
    # Case-insensitive exact match
    for key in patient_data.keys():
        if key.lower() == filename.lower():
            print(f"Found case-insensitive match for '{filename}' with '{key}'")
            return patient_data[key]
    
    # Partial matching (filename contains key or key contains filename)
    for key in patient_data.keys():
        if key.lower() in filename.lower() or filename.lower() in key.lower():
            print(f"Found partial match for '{filename}' with '{key}'")
            return patient_data[key]
    
    # Removing file extension and matching
    filename_no_ext = os.path.splitext(filename)[0]
    for key in patient_data.keys():
        key_no_ext = os.path.splitext(key)[0]
        if key_no_ext.lower() == filename_no_ext.lower():
            print(f"Found match without extension for '{filename}' with '{key}'")
            return patient_data[key]
    
    print(f"No patient data found for '{filename}'")
    return {}

def generate_word_report(prediction_data):
    """Generate Word document report for a prediction"""
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading('Medical Image Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Patient Information
        doc.add_heading('Patient Information', level=1)
        patient_table = doc.add_table(rows=4, cols=2)
        patient_table.style = 'Table Grid'
        
        patient_info = [
            ('Patient ID:', prediction_data.get('patient_id', 'N/A')),
            ('Full Name:', prediction_data.get('patient_name', 'N/A')),
            ('Age:', str(prediction_data.get('patient_age', 'N/A'))),
            ('Symptoms:', prediction_data.get('patient_symptoms', 'N/A'))
        ]
        
        for i, (label, value) in enumerate(patient_info):
            patient_table.cell(i, 0).text = label
            patient_table.cell(i, 1).text = str(value) if value is not None else 'N/A'
        
        # Analysis Results
        doc.add_heading('Analysis Results', level=1)
        results_table = doc.add_table(rows=4, cols=2)
        results_table.style = 'Table Grid'
        
        results_info = [
            ('Image Filename:', prediction_data.get('image_filename', 'N/A')),
            ('Detected Condition:', prediction_data.get('disease', 'N/A')),
            ('Confidence Level:', f"{prediction_data.get('confidence', 0) * 100:.1f}%"),
            ('Analysis Date:', prediction_data.get('timestamp', datetime.utcnow()).strftime('%Y-%m-%d %H:%M:%S'))
        ]
        
        for i, (label, value) in enumerate(results_info):
            results_table.cell(i, 0).text = label
            results_table.cell(i, 1).text = str(value)
        
        # Recommendations
        doc.add_heading('Recommendations', level=1)
        doc.add_paragraph(prediction_data.get('recommendations', 'No specific recommendations available.'))
        
        # Disclaimer
        doc.add_heading('Important Disclaimer', level=1)
        medication = (
            "This is your medication"
        )
        doc.add_paragraph(medication)
        
        return doc
    except Exception as e:
        print(f"Error generating Word report: {e}")
        import traceback
        traceback.print_exc()
        raise


@jwt.expired_token_loader
def expired_token_callback(jwt_header, jwt_payload):
    return jsonify({'error': 'Token has expired'}), 401

@jwt.invalid_token_loader
def invalid_token_callback(error):
    return jsonify({'error': 'Invalid token'}), 401

@jwt.unauthorized_loader
def missing_token_callback(error):
    return jsonify({'error': 'Authorization token is required'}), 401

@app.route('/api/signup', methods=['POST'])
def signup():
    try:
        data = request.get_json()
        if not data.get('name') or not data.get('email') or not data.get('password'):
            return jsonify({'error': 'Missing required fields'}), 400

        if User.query.filter_by(email=data['email']).first():
            return jsonify({'error': 'User already exists'}), 400

        user = User(name=data['name'], email=data['email'])
        user.set_password(data['password'])

        db.session.add(user)
        db.session.commit()

        return jsonify({'message': 'User created successfully'}), 201
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/login', methods=['POST'])
def login():
    try:
        data = request.get_json()
        if not data.get('email') or not data.get('password'):
            return jsonify({'error': 'Missing email or password'}), 400

        user = User.query.filter_by(email=data['email']).first()
        if user and user.check_password(data['password']):
            return jsonify({
                'user': {'id': user.id, 'name': user.name, 'email': user.email}
            }), 200
        return jsonify({'error': 'Invalid credentials'}), 401
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/batch-predict', methods=['POST'])
def batch_predict():
    try:
        user_id = request.form.get('user_id')
        if not user_id:
            return jsonify({'error': 'Missing user_id'}), 400
        
        # Process Excel file if provided and store in global dictionary
        patient_data = {}
        if 'excel_file' in request.files:
            excel_file = request.files['excel_file']
            if excel_file.filename != '':
                print(f"Processing Excel file: {excel_file.filename}")
                patient_data = process_excel_file(excel_file)
                print(f"Patient data extracted: {len(patient_data)} records")
                
                # Store patient data globally for this user
                with patient_data_lock:
                    user_patient_data[user_id] = patient_data
                    print(f"Stored patient data for user {user_id}")
        else:
            # Try to get previously stored patient data for this user
            with patient_data_lock:
                patient_data = user_patient_data.get(user_id, {})
                print(f"Retrieved stored patient data for user {user_id}: {len(patient_data)} records")
        
        # Get all uploaded images
        image_files = request.files.getlist('images')
        if not image_files:
            return jsonify({'error': 'No images uploaded'}), 400
        
        print(f"Processing {len(image_files)} images")
        results = []
        
        for file in image_files:
            if file.filename == '':
                continue
                
            try:
                # Reset file pointer to beginning
                file.seek(0)
                image_bytes = file.read()
                
                print(f"Processing image: {file.filename}")
                prediction_result = predictor.predict(image_bytes)
                
                # Find patient data for this image
                patient_info = find_patient_data(file.filename, patient_data)
                
                print(f"Patient info for {file.filename}: {patient_info}")
                
                prediction = Prediction_table(
                    user_id=user_id,
                    image_data=image_bytes,
                    image_filename=file.filename,
                    disease=prediction_result['disease'],
                    confidence=prediction_result['confidence'],
                    recommendations=prediction_result['recommendations'],
                    patient_id=patient_info.get('id'),
                    patient_name=patient_info.get('full_name'),
                    patient_age=patient_info.get('age'),
                    patient_symptoms=patient_info.get('symptoms')
                )
                db.session.add(prediction)
                db.session.commit()
                
                result = {
                    'id': prediction.id,
                    'fileName': file.filename,
                    'prediction': prediction_result,
                    'patient_data': patient_info if patient_info else None,
                    'success': True
                }
                results.append(result)
                
            except Exception as e:
                print(f"Error processing {file.filename}: {e}")
                import traceback
                traceback.print_exc()
                results.append({
                    'fileName': file.filename,
                    'error': str(e),
                    'success': False
                })
        
        return jsonify({'results': results, 'message': 'Batch prediction completed'}), 200
    except Exception as e:
        print(f"Error in batch_predict: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/predict', methods=['POST'])
def predict():
    try:
        user_id = request.form.get('user_id')
        if 'image' not in request.files:
            return jsonify({'error': 'No image uploaded'}), 400

        file = request.files['image']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        # Reset file pointer to beginning
        file.seek(0)
        image_bytes = file.read()
        prediction_result = predictor.predict(image_bytes)

        if user_id:
            prediction = Prediction_table(
                user_id=user_id,
                image_data=image_bytes,
                image_filename=file.filename,
                disease=prediction_result['disease'],
                confidence=prediction_result['confidence'],
                recommendations=prediction_result['recommendations']
            )
            db.session.add(prediction)
            db.session.commit()

        return jsonify({'prediction': prediction_result, 'message': 'Prediction completed successfully'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/history', methods=['GET'])
def get_history():
    try:
        user_id = request.args.get('user_id')
        if not user_id:
            return jsonify({'error': 'Missing user_id'}), 400
            
        predictions = Prediction_table.query.filter_by(user_id=user_id).order_by(Prediction_table.timestamp.desc()).all()

        history = []
        for p in predictions:
            patient_data = None
            if p.patient_id or p.patient_name or p.patient_age or p.patient_symptoms:
                patient_data = {
                    'id': p.patient_id,
                    'name': p.patient_name,
                    'age': p.patient_age,
                    'symptoms': p.patient_symptoms
                }
            
            history.append({
                'id': p.id,
                'disease': p.disease,
                'confidence': p.confidence,
                'recommendations': p.recommendations,
                'timestamp': p.timestamp.isoformat(),
                'image_filename': p.image_filename,
                'patient_data': patient_data
            })

        return jsonify({'history': history}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/download-report/<int:prediction_id>', methods=['GET'])
def download_report(prediction_id):
    temp_file_path = None
    try:
        user_id = request.args.get('user_id')
        print(f"Download request for prediction_id: {prediction_id}, user_id: {user_id}")
        
        if not user_id:
            return jsonify({'error': 'Missing user_id'}), 400
            
        prediction = Prediction_table.query.filter_by(id=prediction_id, user_id=user_id).first()
        if not prediction:
            print(f"Prediction not found for id: {prediction_id}, user_id: {user_id}")
            return jsonify({'error': 'Prediction not found'}), 404
        
        print(f"Found prediction: {prediction.id}, disease: {prediction.disease}")
        print(f"Patient data: ID={prediction.patient_id}, Name={prediction.patient_name}, Age={prediction.patient_age}")
        
        # Prepare data for Word document
        prediction_data = {
            'patient_id': prediction.patient_id or 'N/A',
            'patient_name': prediction.patient_name or 'N/A',
            'patient_age': prediction.patient_age,
            'patient_symptoms': prediction.patient_symptoms or 'N/A',
            'image_filename': prediction.image_filename or 'N/A',
            'disease': prediction.disease or 'N/A',
            'confidence': prediction.confidence or 0,
            'recommendations': prediction.recommendations or 'N/A',
            'timestamp': prediction.timestamp or datetime.utcnow()
        }
        
        print(f"Prepared prediction data: {prediction_data}")
        
        # Generate Word document
        print("Generating Word document...")
        doc = generate_word_report(prediction_data)
        print("Word document generated successfully")
        
        # Create temporary file with proper cleanup
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx', prefix='medical_report_')
        temp_file_path = temp_file.name
        temp_file.close()
        
        print(f"Created temp file: {temp_file_path}")
        
        # Save document to temporary file
        doc.save(temp_file_path)
        print(f"Document saved to: {temp_file_path}")
        
        # Verify file exists and has content
        if not os.path.exists(temp_file_path):
            raise Exception(f"Temporary file was not created: {temp_file_path}")
        
        file_size = os.path.getsize(temp_file_path)
        print(f"File size: {file_size} bytes")
        
        if file_size == 0:
            raise Exception("Generated file is empty")
        
        # Add to cleanup list
        with cleanup_lock:
            temp_files_to_cleanup.append(temp_file_path)
        
        # Schedule cleanup after 5 minutes
        schedule_cleanup(temp_file_path)
        
        # Generate safe filename
        patient_name = prediction.patient_name or 'Patient'
        # Remove any problematic characters
        safe_patient_name = ''.join(c for c in patient_name if c.isalnum() or c in (' ', '-', '_')).strip()
        if not safe_patient_name:
            safe_patient_name = 'Patient'
        
        filename = f"Medical_Report_{safe_patient_name.replace(' ', '_')}_{prediction.id}.docx"
        print(f"Generated filename: {filename}")
        
        return send_file(
            temp_file_path,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"Error in download_report: {e}")
        import traceback
        traceback.print_exc()
        
        # Clean up temp file if there was an error
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
                print(f"Cleaned up temp file after error: {temp_file_path}")
            except:
                pass
        
        return jsonify({'error': f'Failed to generate report: {str(e)}'}), 500

@app.route('/api/image/<int:prediction_id>', methods=['GET'])
def get_image(prediction_id):
    try:
        user_id = request.args.get('user_id')
        if not user_id:
            return jsonify({'error': 'Missing user_id'}), 400
            
        prediction = Prediction_table.query.filter_by(id=prediction_id, user_id=user_id).first()
        if not prediction:
            return jsonify({'error': 'Prediction not found'}), 404

        image_b64 = base64.b64encode(prediction.image_data).decode('utf-8')
        return jsonify({
            'image_data': image_b64,
            'disease': prediction.disease,
            'confidence': prediction.confidence
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Route to clear patient data
@app.route('/api/clear-patient-data', methods=['POST'])
def clear_patient_data():
    try:
        data = request.get_json()
        user_id = data.get('user_id')
        if not user_id:
            return jsonify({'error': 'Missing user_id'}), 400
        
        with patient_data_lock:
            if user_id in user_patient_data:
                del user_patient_data[user_id]
                print(f"Cleared patient data for user {user_id}")
        
        return jsonify({'message': 'Patient data cleared successfully'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Route to check current patient data
@app.route('/api/patient-data-status', methods=['GET'])
def patient_data_status():
    try:
        user_id = request.args.get('user_id')
        if not user_id:
            return jsonify({'error': 'Missing user_id'}), 400
        
        with patient_data_lock:
            patient_data = user_patient_data.get(user_id, {})
            count = len(patient_data)
            keys = list(patient_data.keys())[:10]  # Show first 10 keys
        
        return jsonify({
            'has_data': count > 0,
            'count': count,
            'sample_keys': keys
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({'status': 'healthy', 'timestamp': datetime.utcnow().isoformat()}), 200

# Register cleanup function to run on exit
atexit.register(cleanup_temp_files)

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True, host='0.0.0.0')